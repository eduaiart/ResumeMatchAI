import os
import logging
import PyPDF2
import docx
from io import BytesIO

class DocumentParser:
    """Parser for different document formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def extract_text(self, file_path):
        """Extract text from a document file"""
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                return self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logging.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
                    
        except Exception as e:
            logging.error(f"Error reading PDF file {file_path}: {str(e)}")
            # Fallback: try with pdfplumber if available
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except ImportError:
                logging.warning("pdfplumber not available, using PyPDF2 result")
            except Exception as e2:
                logging.error(f"Error with pdfplumber fallback: {str(e2)}")
                raise e
        
        return text.strip()
    
    def _extract_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logging.error(f"Error reading DOCX file {file_path}: {str(e)}")
            raise
    
    def _extract_from_txt(self, file_path):
        """Extract text from TXT file"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read().strip()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return file.read().strip()
                
        except Exception as e:
            logging.error(f"Error reading TXT file {file_path}: {str(e)}")
            raise
    
    def is_supported_format(self, filename):
        """Check if file format is supported"""
        file_extension = os.path.splitext(filename)[1].lower()
        return file_extension in self.supported_formats
    
    def get_file_info(self, file_path):
        """Get basic information about the file"""
        try:
            file_stats = os.stat(file_path)
            file_extension = os.path.splitext(file_path)[1].lower()
            
            return {
                'size': file_stats.st_size,
                'format': file_extension,
                'supported': self.is_supported_format(file_path)
            }
        except Exception as e:
            logging.error(f"Error getting file info for {file_path}: {str(e)}")
            return None
